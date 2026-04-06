import re
import zipfile
import os
from io import BytesIO

from PyPDF2 import PdfReader
from docx import Document
import requests

from .ai import build_ai_readiness_snapshot, generate_ai_resume_insights


JOB_ROLES = {
    "frontend-developer": {
        "title": "Frontend Developer",
        "summary": "Builds interactive, accessible, high-performance user interfaces for web products.",
        "skills": [
            "html",
            "css",
            "javascript",
            "typescript",
            "react",
            "redux",
            "tailwind",
            "vite",
            "webpack",
            "responsive design",
            "accessibility",
            "rest api",
            "git",
            "testing",
            "figma",
        ],
        "tools": ["React", "TypeScript", "Figma", "Jest", "Vite"],
        "interview_themes": ["component design", "state management", "browser performance", "accessibility"],
    },
    "backend-developer": {
        "title": "Backend Developer",
        "summary": "Designs APIs, services, data models, and reliable server-side systems.",
        "skills": [
            "python",
            "django",
            "node.js",
            "express",
            "sql",
            "postgresql",
            "mongodb",
            "rest api",
            "authentication",
            "docker",
            "git",
            "testing",
            "redis",
            "microservices",
            "api design",
        ],
        "tools": ["Django", "Express", "PostgreSQL", "Docker", "Redis"],
        "interview_themes": ["api design", "database modeling", "scalability", "security"],
    },
    "data-scientist": {
        "title": "Data Scientist",
        "summary": "Uses data, statistics, and machine learning to generate predictive insights.",
        "skills": [
            "python",
            "pandas",
            "numpy",
            "sql",
            "machine learning",
            "statistics",
            "data visualization",
            "scikit-learn",
            "feature engineering",
            "model evaluation",
            "power bi",
            "tableau",
            "communication",
            "deep learning",
            "nlp",
        ],
        "tools": ["Python", "Pandas", "scikit-learn", "Jupyter", "Tableau"],
        "interview_themes": ["feature engineering", "model metrics", "experimentation", "storytelling"],
    },
    "ml-engineer": {
        "title": "ML Engineer",
        "summary": "Operationalizes machine learning systems and deploys models into production.",
        "skills": [
            "python",
            "machine learning",
            "deep learning",
            "tensorflow",
            "pytorch",
            "mlops",
            "docker",
            "kubernetes",
            "feature engineering",
            "model deployment",
            "sql",
            "fastapi",
            "api design",
            "aws",
            "git",
        ],
        "tools": ["PyTorch", "TensorFlow", "FastAPI", "Docker", "AWS"],
        "interview_themes": ["deployment", "monitoring", "data pipelines", "model tradeoffs"],
    },
    "ui-ux-designer": {
        "title": "UI/UX Designer",
        "summary": "Designs usable, engaging interfaces with strong research and product thinking.",
        "skills": [
            "figma",
            "wireframing",
            "prototyping",
            "user research",
            "design systems",
            "accessibility",
            "visual design",
            "interaction design",
            "usability testing",
            "information architecture",
            "storytelling",
            "responsive design",
            "collaboration",
            "design thinking",
            "journey mapping",
        ],
        "tools": ["Figma", "Framer", "Miro", "Maze", "Adobe XD"],
        "interview_themes": ["design process", "research synthesis", "visual hierarchy", "hand-off quality"],
    },
}

SKILL_SYNONYMS = {
    "js": "javascript",
    "ts": "typescript",
    "node": "node.js",
    "nodejs": "node.js",
    "node.js": "node.js",
    "node-js": "node.js",
    "expressjs": "express",
    "reactjs": "react",
    "react.js": "react",
    "vuejs": "vue",
    "nextjs": "next.js",
    "next.js": "next.js",
    "tailwindcss": "tailwind",
    "tailwind-css": "tailwind",
    "scss": "css",
    "sass": "css",
    "mongodb": "mongodb",
    "mongo": "mongodb",
    "postgres": "postgresql",
    "postgresql": "postgresql",
    "postgre": "postgresql",
    "postgressql": "postgresql",
    "ml": "machine learning",
    "machinelearning": "machine learning",
    "deeplearning": "deep learning",
    "scikitlearn": "scikit-learn",
    "powerbi": "power bi",
    "scikitlearn": "scikit-learn",
    "restful": "rest api",
    "restapis": "rest api",
    "restapi": "rest api",
    "api": "api design",
    "apis": "api design",
    "uiux": "interaction design",
    "wireframes": "wireframing",
    "prototypes": "prototyping",
    "figjam": "figma",
    "nlp": "nlp",
    "aws": "aws",
    "pytorch": "pytorch",
    "tensorflow": "tensorflow",
}

SKILL_ALIASES = {
    "html": ["html", "html5"],
    "css": ["css", "css3", "scss", "sass", "styled components", "styled-components"],
    "javascript": ["javascript", "java script", "js", "ecmascript"],
    "typescript": ["typescript", "type script", "ts", "tsx"],
    "react": ["react", "reactjs", "react.js"],
    "redux": ["redux", "redux toolkit"],
    "tailwind": ["tailwind", "tailwindcss", "tailwind css", "tailwind ui"],
    "vite": ["vite"],
    "webpack": ["webpack", "module bundler"],
    "responsive design": ["responsive design", "responsive ui", "responsive web design", "mobile-first", "mobile first"],
    "accessibility": ["accessibility", "a11y", "wcag", "accessible design", "web accessibility"],
    "rest api": ["rest api", "restful api", "restful apis", "rest services", "api integration"],
    "git": ["git", "github", "gitlab", "version control"],
    "testing": ["testing", "unit testing", "integration testing", "frontend testing", "pytest", "jest", "cypress", "vitest"],
    "figma": ["figma", "figjam"],
    "python": ["python"],
    "django": ["django"],
    "node.js": ["node.js", "nodejs", "node js"],
    "express": ["express", "expressjs", "express js"],
    "sql": ["sql", "mysql", "sqlite"],
    "postgresql": ["postgresql", "postgres", "postgre sql"],
    "mongodb": ["mongodb", "mongo db", "mongo"],
    "authentication": ["authentication", "auth", "oauth", "jwt"],
    "docker": ["docker", "containerization", "containers"],
    "redis": ["redis"],
    "microservices": ["microservices", "micro services"],
    "api design": ["api design", "api architecture", "apis"],
    "pandas": ["pandas"],
    "numpy": ["numpy"],
    "machine learning": ["machine learning", "ml"],
    "statistics": ["statistics", "statistical analysis"],
    "data visualization": ["data visualization", "data visualisation", "visualization", "visualisation"],
    "scikit-learn": ["scikit-learn", "scikit learn", "sklearn"],
    "feature engineering": ["feature engineering"],
    "model evaluation": ["model evaluation", "model validation"],
    "power bi": ["power bi", "powerbi"],
    "tableau": ["tableau"],
    "communication": ["communication", "stakeholder communication"],
    "deep learning": ["deep learning", "deep neural networks"],
    "nlp": ["nlp", "natural language processing"],
    "tensorflow": ["tensorflow", "tensor flow"],
    "pytorch": ["pytorch", "py torch"],
    "mlops": ["mlops", "ml ops", "machine learning operations"],
    "kubernetes": ["kubernetes", "k8s"],
    "model deployment": ["model deployment", "deployment", "serving"],
    "fastapi": ["fastapi", "fast api"],
    "aws": ["aws", "amazon web services"],
    "wireframing": ["wireframing", "wireframes", "wireframe"],
    "prototyping": ["prototyping", "prototype", "prototypes"],
    "user research": ["user research", "ux research"],
    "design systems": ["design systems", "design system"],
    "visual design": ["visual design", "ui design"],
    "interaction design": ["interaction design", "ui/ux", "ui ux", "ux"],
    "usability testing": ["usability testing", "usability tests"],
    "information architecture": ["information architecture", "ia"],
    "storytelling": ["storytelling", "presentation storytelling"],
    "collaboration": ["collaboration", "cross-functional collaboration"],
    "design thinking": ["design thinking"],
    "journey mapping": ["journey mapping", "customer journey mapping"],
}

SECTION_PATTERNS = {
    "summary": ["summary", "profile", "objective", "about"],
    "skills": ["skills", "technical skills", "core skills", "tech stack"],
    "projects": ["projects", "project experience", "personal projects"],
    "experience": ["experience", "work experience", "employment", "internship"],
    "education": ["education", "academic background", "qualification"],
}


LEARNING_RESOURCES = {
    "html": [
        {"title": "HTML roadmap", "url": "https://roadmap.sh/html", "action": "Follow the full HTML learning roadmap step by step."},
        {"title": "MDN HTML course", "url": "https://developer.mozilla.org/en-US/docs/Learn_web_development/Core/Structuring_content", "action": "Work through the MDN HTML learning path with examples."},
    ],
    "css": [
        {"title": "CSS roadmap", "url": "https://roadmap.sh/css", "action": "Use the CSS roadmap to cover layout, selectors, and advanced topics."},
        {"title": "MDN CSS guide", "url": "https://developer.mozilla.org/en-US/docs/Web/CSS", "action": "Learn CSS fundamentals and reference real properties quickly."},
    ],
    "javascript": [
        {"title": "JavaScript roadmap", "url": "https://roadmap.sh/javascript", "action": "Follow the JavaScript roadmap in the correct order."},
        {"title": "javascript.info course", "url": "https://javascript.info/", "action": "Study JavaScript from basics to advanced topics with examples."},
    ],
    "typescript": [
        {"title": "TypeScript roadmap", "url": "https://roadmap.sh/typescript", "action": "Follow the TypeScript roadmap from basics to advanced usage."},
        {"title": "TypeScript handbook", "url": "https://www.typescriptlang.org/docs/", "action": "Learn TypeScript directly from the official handbook."},
    ],
    "react": [
        {"title": "React roadmap", "url": "https://roadmap.sh/react", "action": "Use the React roadmap to learn the ecosystem in order."},
        {"title": "React official course", "url": "https://react.dev/learn", "action": "Study components, state, effects, and React patterns."},
    ],
    "redux": [
        {"title": "Redux essentials tutorial", "url": "https://redux.js.org/tutorials/essentials/part-1-overview-concepts", "action": "Learn Redux with the official end-to-end tutorial."},
        {"title": "Redux Toolkit quick start", "url": "https://redux-toolkit.js.org/tutorials/quick-start", "action": "Set up and use the modern Redux Toolkit workflow."},
    ],
    "tailwind": [
        {"title": "Tailwind CSS docs", "url": "https://tailwindcss.com/docs", "action": "Learn utility classes, styling patterns, and customization."},
        {"title": "Tailwind installation guide", "url": "https://tailwindcss.com/docs/installation", "action": "Set up Tailwind correctly in a project and start practicing."},
    ],
    "vite": [
        {"title": "Vite official guide", "url": "https://vite.dev/guide/", "action": "Learn how to scaffold, run, and build apps with Vite."},
        {"title": "Vite features guide", "url": "https://vite.dev/guide/features.html", "action": "Understand assets, env files, and the dev workflow."},
    ],
    "responsive design": [
        {"title": "Responsive design course", "url": "https://web.dev/learn/design", "action": "Study breakpoints, fluid layouts, and mobile-first design."},
        {"title": "MDN responsive design", "url": "https://developer.mozilla.org/en-US/docs/Learn_web_development/Core/CSS_layout/Responsive_Design", "action": "Practice media queries and responsive CSS techniques."},
    ],
    "accessibility": [
        {"title": "W3C accessibility introduction", "url": "https://www.w3.org/WAI/fundamentals/accessibility-intro/", "action": "Learn core accessibility principles and why they matter."},
        {"title": "MDN accessibility guide", "url": "https://developer.mozilla.org/en-US/docs/Web/Accessibility", "action": "Apply keyboard, semantic, and assistive-technology-friendly practices."},
    ],
    "rest api": [
        {"title": "REST API tutorial", "url": "https://restfulapi.net/", "action": "Learn REST methods, status codes, and API design basics."},
        {"title": "MDN REST guide", "url": "https://developer.mozilla.org/en-US/docs/Glossary/REST", "action": "Understand what REST means and how web APIs are designed."},
    ],
    "git": [
        {"title": "Pro Git book", "url": "https://git-scm.com/book/en/v2", "action": "Learn commits, branches, merges, and collaboration workflows."},
        {"title": "Git documentation", "url": "https://git-scm.com/docs", "action": "Use the official documentation for exact commands and concepts."},
    ],
    "testing": [
        {"title": "Jest getting started", "url": "https://jestjs.io/docs/getting-started", "action": "Learn unit testing and configure a frontend test setup."},
        {"title": "Testing Library docs", "url": "https://testing-library.com/docs/", "action": "Practice user-focused testing patterns used in real apps."},
    ],
    "figma": [
        {"title": "Figma design basics", "url": "https://www.figma.com/resource-library/design-basics/", "action": "Learn UI basics, layouts, and design workflow in Figma."},
        {"title": "Figma help center", "url": "https://help.figma.com/hc/en-us", "action": "Use the official Figma help center to learn tools and features."},
    ],
    "python": [
        {"title": "Python roadmap", "url": "https://roadmap.sh/python", "action": "Follow the roadmap for learning Python in the right sequence."},
        {"title": "Python docs tutorial", "url": "https://docs.python.org/3/tutorial/", "action": "Learn Python directly from the official tutorial."},
    ],
    "django": [
        {"title": "Django roadmap", "url": "https://roadmap.sh/django", "action": "Use the Django roadmap to understand the backend stack."},
        {"title": "Django tutorial", "url": "https://docs.djangoproject.com/en/stable/intro/tutorial01/", "action": "Build a Django app and learn views, models, and routing."},
    ],
    "postgresql": [
        {"title": "PostgreSQL roadmap", "url": "https://roadmap.sh/postgresql", "action": "Follow a guided roadmap for PostgreSQL concepts and usage."},
        {"title": "PostgreSQL tutorial", "url": "https://www.postgresql.org/docs/current/tutorial.html", "action": "Learn SQL queries and PostgreSQL basics from the official docs."},
    ],
    "docker": [
        {"title": "Docker roadmap", "url": "https://roadmap.sh/docker", "action": "Use the Docker roadmap to learn containers in sequence."},
        {"title": "Docker getting started", "url": "https://docs.docker.com/get-started/", "action": "Learn images, containers, and the Docker workflow."},
    ],
    "machine learning": [
        {"title": "Machine learning roadmap", "url": "https://roadmap.sh/ai-data-scientist", "action": "Use the AI and data roadmap to structure ML learning."},
        {"title": "Google ML crash course", "url": "https://developers.google.com/machine-learning/crash-course", "action": "Study ML fundamentals with Google's guided course."},
    ],
    "scikit-learn": [
        {"title": "scikit-learn user guide", "url": "https://scikit-learn.org/stable/user_guide.html", "action": "Learn preprocessing, models, and evaluation in sklearn."},
        {"title": "scikit-learn tutorials", "url": "https://scikit-learn.org/stable/tutorial/index.html", "action": "Work through example-driven sklearn tutorials."},
    ],
    "tensorflow": [
        {"title": "TensorFlow tutorials", "url": "https://www.tensorflow.org/tutorials", "action": "Learn TensorFlow using guided official tutorials."},
        {"title": "TensorFlow guide", "url": "https://www.tensorflow.org/guide", "action": "Study the main TensorFlow workflow in depth."},
    ],
    "pytorch": [
        {"title": "PyTorch tutorials", "url": "https://pytorch.org/tutorials/", "action": "Learn tensors, training loops, and model building."},
        {"title": "Learn PyTorch", "url": "https://www.learnpytorch.io/", "action": "Follow a practical end-to-end PyTorch learning path."},
    ],
    "fastapi": [
        {"title": "FastAPI tutorial", "url": "https://fastapi.tiangolo.com/tutorial/", "action": "Learn API building with the official FastAPI tutorial."},
        {"title": "FastAPI docs", "url": "https://fastapi.tiangolo.com/", "action": "Use the official reference while building backend projects."},
    ],
}

YOUTUBE_COURSES = {
    "html": [
        {
            "title": "HTML & CSS full beginner course",
            "url": "https://www.youtube.com/watch?v=G3e-cpL7ofc",
            "action": "Watch the HTML sections first, then rebuild one semantic landing page with forms and headings.",
        },
    ],
    "css": [
        {
            "title": "HTML & CSS full beginner course",
            "url": "https://www.youtube.com/watch?v=G3e-cpL7ofc",
            "action": "Use the CSS sections to practice spacing, layout, typography, and responsive structure.",
        },
    ],
    "javascript": [
        {
            "title": "JavaScript full beginner course",
            "url": "https://www.youtube.com/watch?v=EerdGm-ehJQ",
            "action": "Follow the full JavaScript course and build one interactive project while learning.",
        },
    ],
    "typescript": [
        {
            "title": "Practical TypeScript for beginners",
            "url": "https://www.youtube.com/watch?v=JHEB7RhJG1Y",
            "action": "Complete the TypeScript course and convert one JavaScript app to typed code.",
        },
    ],
    "react": [
        {
            "title": "React course for beginners",
            "url": "https://www.youtube.com/watch?v=ZXIN4i_xJ6Q",
            "action": "Finish the React course and ship a component-based mini app.",
        },
    ],
    "redux": [
        {
            "title": "Redux Toolkit playlist",
            "url": "https://www.youtube.com/playlist?list=PL0Zuz27SZ-6M1J5I1w2-uZx36Qp6qhjKo",
            "action": "Work through the Redux Toolkit playlist and wire shared state into a React app.",
        },
    ],
    "git": [
        {
            "title": "Git and GitHub full course",
            "url": "https://www.youtube.com/watch?v=RGOj5yH7evk",
            "action": "Learn commits, branches, merges, and GitHub workflow with the full beginner course.",
        },
    ],
    "tailwind": [
        {
            "title": "Tailwind CSS tutorial series",
            "url": "https://www.youtube.com/@NetNinja/playlists",
            "action": "Open the Net Ninja playlists page and complete the Tailwind CSS tutorial series.",
        },
    ],
    "vite": [
        {
            "title": "Vite beginner tutorial path",
            "url": "https://www.youtube.com/@HiteshCodeLab/playlists",
            "action": "Use the Hitesh Choudhary playlists to complete a Vite setup tutorial and ship one build.",
        },
    ],
    "responsive design": [
        {
            "title": "Responsive web design course",
            "url": "https://www.youtube.com/@freecodecamp/videos",
            "action": "Open freeCodeCamp's channel and complete one responsive web design course or long tutorial.",
        },
    ],
    "accessibility": [
        {
            "title": "Web accessibility learning playlist",
            "url": "https://www.youtube.com/@GoogleChromeDevelopers/videos",
            "action": "Use the Chrome Developers channel and complete the accessibility-focused lessons before auditing your UI.",
        },
    ],
    "rest api": [
        {
            "title": "REST API beginner course",
            "url": "https://www.youtube.com/@HiteshCodeLab/videos",
            "action": "Complete one beginner REST API tutorial and connect a public API inside a small project.",
        },
    ],
    "testing": [
        {
            "title": "Frontend testing tutorial path",
            "url": "https://www.youtube.com/@freecodecamp/videos",
            "action": "Pick one Jest or Testing Library beginner course from freeCodeCamp and add tests to a project.",
        },
    ],
    "figma": [
        {
            "title": "Figma beginner playlist",
            "url": "https://www.youtube.com/results?search_query=figma+for+beginners+freecodecamp",
            "action": "Complete one beginner-friendly Figma course and recreate a polished UI screen.",
        },
    ],
}


def normalize_whitespace(text):
    return re.sub(r"\s+", " ", text).strip()


def repair_fragmented_pdf_text(text):
    repaired = text or ""
    repaired = re.sub(r"(?<=\w)-\s+(?=\w)", "", repaired)
    repaired = re.sub(r"(?<=\w)\s{2,}(?=\w)", " ", repaired)

    leading_split_word = re.compile(r"\b(?:[A-Za-z]\s){1,}[A-Za-z][A-Za-z0-9+#.-]*\b")
    repaired = leading_split_word.sub(lambda match: match.group(0).replace(" ", ""), repaired)

    fully_split_word = re.compile(r"\b(?:[A-Za-z0-9+#.-]\s){2,}[A-Za-z0-9+#.-]\b")
    repaired = fully_split_word.sub(lambda match: match.group(0).replace(" ", ""), repaired)
    return repaired


def build_search_variants(text):
    raw_text = text or ""
    repaired_text = repair_fragmented_pdf_text(raw_text)
    variants = []

    for candidate in (raw_text, repaired_text):
        lowered = candidate.lower()
        variants.extend(
            [
                lowered,
                normalize_resume_text(candidate),
                re.sub(r"[^a-z0-9]+", "", lowered),
            ]
        )

    unique_variants = []
    seen = set()
    for item in variants:
        compact = item.strip()
        if compact and compact not in seen:
            unique_variants.append(compact)
            seen.add(compact)
    return unique_variants


def extract_resume_text(uploaded_file):
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()

    if name.endswith(".txt"):
        return repair_fragmented_pdf_text(raw.decode("utf-8", errors="ignore"))

    if name.endswith(".pdf"):
        reader = PdfReader(BytesIO(raw))
        text = repair_fragmented_pdf_text("\n".join((page.extract_text() or "") for page in reader.pages).strip())
        if not has_meaningful_text(text):
            if is_likely_scanned_pdf(raw):
                raise ValueError(
                    "This PDF looks like a scanned or image-based resume and cannot be analyzed right now. "
                    "Please upload a text-based PDF, DOCX, or TXT file instead."
                )
            raise ValueError(
                "Text could not be extracted reliably from this PDF. "
                "Please try a DOCX or TXT file."
            )
        return text

    if name.endswith(".docx"):
        try:
            document = Document(BytesIO(raw))
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)
            return repair_fragmented_pdf_text("\n".join(parts).strip())
        except Exception:
            return repair_fragmented_pdf_text(extract_docx_text_fallback(raw))

    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")


def normalize_resume_text(text):
    normalized = text.lower()
    normalized = normalized.replace("&", " and ")
    normalized = re.sub(r"(?<=[a-z])/(?=[a-z])", " ", normalized)
    normalized = re.sub(r"(?<=[a-z])-(?=[a-z])", " ", normalized)
    normalized = re.sub(r"(?<=[a-z])\.(?=[a-z])", "", normalized)
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def build_loose_skill_pattern(variant):
    cleaned = variant.lower().strip()
    parts = re.split(r"[\s./+-]+", cleaned)
    token_patterns = []

    for part in parts:
        alnum = re.sub(r"[^a-z0-9]", "", part)
        if not alnum:
            continue
        token_patterns.append(r"\s*".join(re.escape(char) for char in alnum))

    if not token_patterns:
        return None

    joined = r"[\s./+-]*".join(token_patterns)
    return rf"(^|[^a-z0-9]){joined}([^a-z0-9]|$)"


def has_meaningful_text(text):
    cleaned = re.sub(r"\s+", " ", text).strip()
    return len(cleaned) >= 40 and len(re.findall(r"[a-zA-Z]{2,}", cleaned)) >= 8


def is_likely_scanned_pdf(raw):
    pdf_text = raw.decode("latin1", errors="ignore")
    has_image_object = "/Subtype /Image" in pdf_text or "/XObject" in pdf_text
    has_text_font = "/Font" in pdf_text or "BT" in pdf_text
    return has_image_object and not has_meaningful_text(pdf_text) and not has_text_font


def extract_text_with_ocr(raw, filename):
    api_key = os.getenv("OCR_SPACE_API_KEY", "").strip()
    if not api_key:
        return ""

    try:
        response = requests.post(
            "https://api.ocr.space/parse/image",
            data={
                "apikey": api_key,
                "language": "eng",
                "isOverlayRequired": False,
                "detectOrientation": True,
                "scale": True,
                "OCREngine": 2,
            },
            files={"file": (filename, raw)},
            timeout=60,
        )
        response.raise_for_status()
        payload = response.json()
    except Exception:
        return ""

    parsed_results = payload.get("ParsedResults") or []
    texts = [item.get("ParsedText", "") for item in parsed_results if item.get("ParsedText")]
    return "\n".join(texts).strip()


def extract_docx_text_fallback(raw):
    with zipfile.ZipFile(BytesIO(raw)) as docx_zip:
        xml = docx_zip.read("word/document.xml").decode("utf-8", errors="ignore")
    text = re.sub(r"<w:p[^>]*>", "\n", xml)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def analyze_resume(text, role_key, filename, job_description="", raw_file_bytes=None):
    role = JOB_ROLES[role_key]
    extracted_skills = extract_skills(text)
    jd_analysis = analyze_job_description(job_description, role)
    target_skills = jd_analysis["target_skills"] or role["skills"]
    matched_skills = [skill for skill in target_skills if skill in extracted_skills]
    missing_skills = [skill for skill in target_skills if skill not in extracted_skills]
    match_rate = round((len(matched_skills) / len(target_skills)) * 100) if target_skills else 0
    ats = compute_ats_signals(text)
    resume_score = round(match_rate * 0.6 + ats["score"] * 0.4)
    experience_level = detect_experience_level(text)
    highlights = extract_highlights(text)
    section_scores = compute_section_scores(text, extracted_skills)
    prioritized_gaps = prioritize_missing_skills(missing_skills, jd_analysis["jd_skills"])
    course_recommendations = build_course_recommendations(missing_skills)
    bullet_improvements = build_bullet_improvements(highlights, matched_skills, missing_skills)
    recommendations = build_recommendations(missing_skills)
    learning_roadmap = build_learning_roadmap(missing_skills)
    interview_questions = build_interview_questions(role, matched_skills, missing_skills)
    strengths = [
        f"{len(matched_skills)} relevant skills match the target role.",
        *ats["strengths"][:2],
        "Resume includes readable content blocks that can be sharpened into stronger bullets." if highlights else "",
    ]
    warnings = ats["warnings"]
    ai_readiness = build_ai_readiness_snapshot(
        filename=filename,
        role_title=role["title"],
        role_summary=role["summary"],
        resume_text=text,
        job_description=job_description,
        matched_skills=matched_skills,
        missing_skills=missing_skills,
        resume_score=resume_score,
        match_rate=match_rate,
        ats_score=ats["score"],
        roadmap=learning_roadmap,
    )
    ai_result = generate_ai_resume_insights(
        filename=filename,
        role_title=role["title"],
        role_summary=role["summary"],
        extracted_text=text,
        job_description=job_description,
        matched_skills=matched_skills,
        missing_skills=missing_skills,
        rule_score=resume_score,
        match_rate=match_rate,
        ats_score=ats["score"],
        raw_file_bytes=raw_file_bytes,
    )

    if ai_result["active"] and ai_result["ai_score"] is not None:
        resume_score = round(resume_score * 0.55 + ai_result["ai_score"] * 0.45)
        if ai_result["roadmap"]:
            learning_roadmap = ai_result["roadmap"]
        if ai_result["recommendations"]:
            merged_recommendations = []
            for item in [*ai_result["recommendations"], *recommendations]:
                cleaned = normalize_whitespace(item)
                if cleaned and cleaned not in merged_recommendations:
                    merged_recommendations.append(cleaned)
            recommendations = merged_recommendations
        if ai_result["strengths"]:
            for item in ai_result["strengths"]:
                cleaned = normalize_whitespace(item)
                if cleaned and cleaned not in strengths:
                    strengths.append(cleaned)
        if ai_result["gaps"]:
            for item in ai_result["gaps"]:
                cleaned = normalize_whitespace(item)
                if cleaned and cleaned not in warnings:
                    warnings.append(cleaned)

        ai_readiness["active"] = True
        ai_readiness["status"] = "active"
        ai_readiness["provider"] = ai_result["provider"]
        ai_readiness["model"] = ai_result["model"]

    analysis = {
        "filename": filename,
        "role": role["title"],
        "role_summary": role["summary"],
        "job_description_used": bool(job_description.strip()),
        "job_description_summary": jd_analysis["summary"],
        "resume_score": resume_score,
        "rule_based_resume_score": round(match_rate * 0.6 + ats["score"] * 0.4),
        "score_mode": "Hybrid AI + Rule-based" if ai_result["active"] and ai_result["ai_score"] is not None else "Rule-based",
        "match_rate": match_rate,
        "ats_score": ats["score"],
        "experience_level": experience_level,
        "matched_skills": matched_skills,
        "all_missing_skills": missing_skills,
        "missing_skills": missing_skills,
        "completed_skills": [],
        "prioritized_gaps": prioritized_gaps,
        "section_scores": section_scores,
        "course_recommendations": course_recommendations,
        "bullet_improvements": bullet_improvements,
        "strengths": strengths,
        "warnings": warnings,
        "recommendations": recommendations,
        "learning_roadmap": learning_roadmap,
        "interview_questions": interview_questions,
        "highlights": highlights,
        "suggested_tools": role["tools"],
        "ai_readiness": ai_readiness,
        "ai_score": ai_result["ai_score"],
        "ai_summary": ai_result["summary"],
        "ai_learning_roadmap": ai_result["roadmap"],
        "ai_error": ai_result["error"],
        "jd_skills": jd_analysis["jd_skills"],
        "base_resume_score": resume_score,
        "base_match_rate": match_rate,
    }
    return apply_learning_progress(analysis)


def tokenize_text(text):
    cleaned = re.sub(r"[^\w.+#/-]+", " ", text.lower())
    return [token for token in cleaned.split() if token]


def normalize_skill(token):
    normalized = re.sub(r"[^a-z0-9]", "", token.lower())
    return SKILL_SYNONYMS.get(normalized, token.lower())


def extract_skills(text):
    search_variants = build_search_variants(text)
    known_skills = {skill for role in JOB_ROLES.values() for skill in role["skills"]}
    found = set()

    for skill in known_skills:
        variants = list(dict.fromkeys([skill, *SKILL_ALIASES.get(skill, [])]))
        for variant in variants:
            loose_pattern = build_loose_skill_pattern(variant)
            normalized_variant = re.sub(r"[^a-z0-9]+", "", variant.lower())
            exact_pattern = rf"(^|[^a-z0-9]){re.escape(variant.lower())}([^a-z0-9]|$)"

            for search_text in search_variants:
                compressed_search = re.sub(r"[^a-z0-9]+", "", search_text)
                compressed_match = normalized_variant and len(normalized_variant) >= 4 and normalized_variant in compressed_search
                if (
                    re.search(exact_pattern, search_text)
                    or (loose_pattern and re.search(loose_pattern, search_text))
                    or compressed_match
                ):
                    found.add(skill)
                    break

            if skill in found:
                break

    for token in tokenize_text(text):
        normalized = normalize_skill(token)
        if normalized in known_skills:
            found.add(normalized)

    return sorted(found)


def analyze_job_description(job_description, role):
    if not job_description or not job_description.strip():
        return {
            "summary": "Using the built-in role profile because no custom job description was provided.",
            "jd_skills": [],
            "target_skills": role["skills"],
        }

    jd_skills = extract_skills(job_description)
    target_skills = []
    for skill in role["skills"]:
        if skill in jd_skills or not jd_skills:
            target_skills.append(skill)
    for skill in jd_skills:
        if skill not in target_skills:
            target_skills.append(skill)

    return {
        "summary": f"Compared against a custom job description with {len(jd_skills)} detected relevant skills.",
        "jd_skills": jd_skills,
        "target_skills": target_skills or role["skills"],
    }


def extract_section_text(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    sections = {name: [] for name in SECTION_PATTERNS}
    current_section = None

    for line in lines:
        matched_section = None
        line_lower = line.lower()
        for section, markers in SECTION_PATTERNS.items():
            if any(line_lower == marker or line_lower.startswith(f"{marker}:") for marker in markers):
                matched_section = section
                break
        if matched_section:
            current_section = matched_section
            continue
        if current_section:
            sections[current_section].append(line)

    return {name: "\n".join(content).strip() for name, content in sections.items()}


def compute_section_scores(text, extracted_skills):
    sections = extract_section_text(text)
    score_map = []
    for section, content in sections.items():
        if not content:
            score = 18 if section != "summary" else 10
        else:
            section_skills = extract_skills(content)
            length_bonus = min(len(content.split()), 80) // 4
            score = min(100, 25 + len(section_skills) * 12 + length_bonus)
        score_map.append({"section": section.title(), "score": score})
    return score_map


def prioritize_missing_skills(missing_skills, jd_skills):
    priorities = []
    for index, skill in enumerate(missing_skills):
        if skill in jd_skills or index < 2:
            label = "High"
        elif index < 5:
            label = "Medium"
        else:
            label = "Low"
        priorities.append({"skill": skill, "priority": label})
    return priorities


def build_course_recommendations(missing_skills):
    items = []
    for skill in missing_skills[:5]:
        resources = YOUTUBE_COURSES.get(skill) or LEARNING_RESOURCES.get(
            skill,
            [
                {
                    "title": f"{skill.title()} learning path",
                    "url": f"https://www.youtube.com/results?search_query={skill.replace(' ', '+')}+full+course",
                    "action": f"Open a structured beginner-friendly video path for {skill} and complete one hands-on project.",
                },
                {
                    "title": f"{skill.title()} official guide",
                    "url": "https://developer.mozilla.org/",
                    "action": f"Use the official documentation path for {skill}.",
                },
            ],
        )
        items.append({"skill": skill, "resources": resources})
    return items


def apply_learning_progress(analysis):
    completed_skills = sorted({skill for skill in analysis.get("completed_skills", []) if skill})
    original_missing = analysis.get("all_missing_skills") or list(
        dict.fromkeys([*(analysis.get("missing_skills") or []), *completed_skills])
    )
    matched_skills = analysis.get("matched_skills", [])
    remaining_missing = [skill for skill in original_missing if skill not in completed_skills]
    completed_relevant = [skill for skill in completed_skills if skill in original_missing]
    total_target_skills = len(matched_skills) + len(original_missing)

    projected_match_rate = analysis.get("base_match_rate", analysis.get("match_rate", 0))
    if total_target_skills:
        projected_match_rate = round(((len(matched_skills) + len(completed_relevant)) / total_target_skills) * 100)

    projected_rule_score = round(projected_match_rate * 0.6 + analysis.get("ats_score", 0) * 0.4)
    ai_score = analysis.get("ai_score")
    if ai_score is not None:
        projected_resume_score = max(
            analysis.get("base_resume_score", analysis.get("resume_score", 0)),
            round(projected_rule_score * 0.55 + ai_score * 0.45),
        )
    else:
        projected_resume_score = max(
            analysis.get("base_resume_score", analysis.get("resume_score", 0)),
            projected_rule_score,
        )

    ai_learning_roadmap = analysis.get("ai_learning_roadmap") or []
    roadmap = []
    if ai_learning_roadmap:
        roadmap = [item for item in ai_learning_roadmap if item.get("focus") not in completed_relevant]
    if not roadmap:
        roadmap = build_learning_roadmap(remaining_missing)

    analysis["completed_skills"] = completed_relevant
    analysis["all_missing_skills"] = original_missing
    analysis["missing_skills"] = remaining_missing
    analysis["prioritized_gaps"] = prioritize_missing_skills(remaining_missing, analysis.get("jd_skills", []))
    analysis["learning_roadmap"] = roadmap
    analysis["course_recommendations"] = build_course_recommendations(remaining_missing)
    analysis["match_rate"] = projected_match_rate
    analysis["resume_score"] = projected_resume_score
    analysis["score_gain"] = projected_resume_score - analysis.get("base_resume_score", projected_resume_score)
    analysis["match_gain"] = projected_match_rate - analysis.get("base_match_rate", projected_match_rate)
    analysis["score_mode"] = analysis.get("score_mode", "Rule-based")
    return analysis


def build_bullet_improvements(highlights, matched_skills, missing_skills):
    suggestions = []
    metric_focus = missing_skills[0] if missing_skills else "business impact"

    candidate_lines = []
    for line in highlights:
        cleaned = normalize_whitespace(line)
        lower = cleaned.lower()
        if len(cleaned.split()) < 7:
            continue
        if not any(
            verb in lower
            for verb in [
                "built",
                "develop",
                "improv",
                "design",
                "lead",
                "implement",
                "create",
                "migrat",
                "optimiz",
                "integrat",
                "launch",
                "deliver",
                "refactor",
                "manage",
                "collaborat",
                "architect",
            ]
        ):
            continue
        if lower.count("github.com") or lower.count("@") or re.search(r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b", lower):
            continue
        if re.search(r"\b(?:present|owner|developer|contributor|stars?)\b", lower) and not any(char in lower for char in ["%", "improv", "built", "integrat", "deliver", "optimiz", "refactor"]):
            continue
        candidate_lines.append(cleaned)

    def display_skill(skill):
        special_labels = {
            "api design": "API design",
            "rest api": "REST API",
            "ui/ux": "UI/UX",
        }
        return special_labels.get(skill.lower(), skill.title())

    def rewrite_line(line, index):
        lower = line.lower()
        line_skills = extract_skills(line)
        base_skill = line_skills[0] if line_skills else (matched_skills[0] if matched_skills else "core stack")
        action_skill = display_skill(base_skill)
        metric_phrase = "improved load time by 30%" if index == 0 else "reduced manual effort and improved delivery quality"

        if "built" in lower or "develop" in lower or "create" in lower:
            return f"Built and shipped {action_skill}-based features that supported user needs and {metric_phrase}, while documenting clear ownership and outcomes."
        if "improv" in lower or "optimiz" in lower or "performance" in lower:
            return f"Optimized the application using {action_skill} practices, {metric_phrase}, and improved the overall user experience."
        if "migrat" in lower or "refactor" in lower:
            return f"Refactored and migrated existing features with {action_skill}, making the codebase easier to maintain and creating room to strengthen {metric_focus} coverage."
        if "design" in lower or "ui" in lower or "ux" in lower:
            return f"Designed and delivered user-facing interfaces with stronger {action_skill} execution, clearer collaboration, and measurable improvements in usability."
        return f"Delivered work using {action_skill} that created visible product value, highlighted ownership, and should be rewritten with metrics tied to {metric_focus}."

    for index, line in enumerate(candidate_lines[:3]):
        suggestions.append({"before": line, "after": rewrite_line(line, index)})

    if not suggestions:
        return []

    return suggestions


def compute_ats_signals(text):
    lower = text.lower()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    bullet_count = len(re.findall(r"(?:\u2022|-)\s", text))
    quantified_achievements = len(re.findall(r"\b\d+%|\b\d+\+|\$\d+", text))
    email_present = bool(re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text, re.I))
    phone_present = bool(re.search(r"(\+\d{1,3}[-.\s]?)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}", text))
    linkedin_present = "linkedin.com" in lower
    github_present = "github.com" in lower
    section_keywords = ["summary", "experience", "education", "skills", "projects"]
    present_sections = [section for section in section_keywords if section in lower]
    warnings = []

    if not email_present:
        warnings.append("Add a professional email address.")
    if not phone_present:
        warnings.append("Include a contact phone number.")
    if not linkedin_present:
        warnings.append("Add a LinkedIn profile for recruiter trust.")
    if bullet_count < 4:
        warnings.append("Use more bullet points to improve scanability.")
    if quantified_achievements < 2:
        warnings.append("Add measurable outcomes like percentages, counts, or revenue impact.")
    if len(present_sections) < 4:
        warnings.append("Structure the resume with clear ATS-friendly sections.")
    if any(len(line) > 140 for line in lines):
        warnings.append("Shorten long lines so key details are not buried.")

    score = max(
        38,
        min(
            96,
            48
            + len(present_sections) * 7
            + min(bullet_count, 10) * 2
            + quantified_achievements * 5
            + (6 if email_present else 0)
            + (4 if phone_present else 0)
            + (4 if linkedin_present else 0)
            + (3 if github_present else 0)
            - len(warnings) * 4,
        ),
    )

    strengths = []
    if email_present:
        strengths.append("Contains an email address.")
    if phone_present:
        strengths.append("Contains a phone number.")
    if linkedin_present:
        strengths.append("Includes LinkedIn profile.")
    if quantified_achievements >= 2:
        strengths.append("Shows measurable impact.")
    if len(present_sections) >= 4:
        strengths.append("Uses standard ATS-friendly sections.")

    return {"score": score, "warnings": warnings, "strengths": strengths}


def detect_experience_level(text):
    lower = text.lower()
    years = [int(match.group(1)) for match in re.finditer(r"(\d+)\+?\s+years?", lower)]
    max_years = max(years) if years else 0
    if max_years >= 5 or re.search(r"senior|lead|architect|principal", lower):
        return "Advanced"
    if max_years >= 2 or re.search(r"internship|associate|junior", lower):
        return "Intermediate"
    return "Early career"


def extract_highlights(text):
    lines = [line.strip() for line in text.splitlines()]
    return [line for line in lines if 25 <= len(line) <= 150][:4]


def build_recommendations(missing_skills):
    recommendations = [
        "Rewrite at least two bullets using action verbs plus measurable outcomes.",
        "Tailor the summary section to the chosen role instead of keeping it generic.",
    ]
    if missing_skills:
        recommendations.insert(0, f"Prioritize {missing_skills[0]} because it appears in the target role profile.")
    if len(missing_skills) > 1:
        recommendations.insert(1, f"Add one project bullet demonstrating {missing_skills[1]}.")
    return recommendations


ROADMAP_PLAYBOOK = {
    "html": {
        "mission": "Learn semantic HTML, forms, and accessibility basics, then rebuild one landing page using clean sectioning tags.",
        "output": "A single-page semantic website with header, main, sections, form, and accessible labels pushed to GitHub.",
    },
    "css": {
        "mission": "Practice layout, spacing, typography, and responsive styling by recreating a polished dashboard screen from scratch.",
        "output": "A responsive UI clone that demonstrates Flexbox, Grid, spacing, and modern CSS structure.",
    },
    "javascript": {
        "mission": "Strengthen DOM, events, arrays, and async logic by building an interactive app with real user actions.",
        "output": "A small JavaScript app such as task manager, weather app, or quiz with clean event handling.",
    },
    "typescript": {
        "mission": "Learn types, interfaces, props typing, and API typing by converting one JavaScript project into TypeScript.",
        "output": "A TypeScript project that uses interfaces, typed components or functions, and zero type errors.",
    },
    "react": {
        "mission": "Build reusable components, state handling, and props flow through a multi-section frontend project.",
        "output": "A React app with at least 5 reusable components, stateful UI, and deployed demo link.",
    },
    "redux": {
        "mission": "Practice centralized state management by adding Redux Toolkit to a project with shared app state.",
        "output": "A React project using Redux slices, store setup, and one async or shared-state workflow.",
    },
    "tailwind": {
        "mission": "Learn utility-first styling by converting one existing page into Tailwind with responsive breakpoints.",
        "output": "A Tailwind-based UI with reusable utility patterns, responsive sections, and polished spacing.",
    },
    "vite": {
        "mission": "Understand modern frontend tooling by creating a project in Vite and managing assets, scripts, and environment setup.",
        "output": "A Vite-based app scaffolded, run locally, built successfully, and documented in the README.",
    },
    "webpack": {
        "mission": "Learn bundling concepts by setting up entry, output, loaders, and build optimization in a simple project.",
        "output": "A small project with a working Webpack config, asset handling, and production build command.",
    },
    "responsive design": {
        "mission": "Design mobile-first layouts and test them at common breakpoints using one realistic product page.",
        "output": "A page that works cleanly on mobile, tablet, and desktop with documented breakpoint decisions.",
    },
    "accessibility": {
        "mission": "Audit one page for keyboard access, labels, color contrast, and semantic structure, then fix the issues.",
        "output": "An accessibility-improved UI with visible focus states, semantic tags, and a short audit checklist.",
    },
    "rest api": {
        "mission": "Learn request methods, endpoints, and API integration by consuming one public REST API in a project.",
        "output": "A frontend or backend demo that fetches, displays, and handles API responses or errors correctly.",
    },
    "git": {
        "mission": "Practice commits, branches, merges, and pull-style workflows on a small project repository.",
        "output": "A GitHub repo with clean commit history, one feature branch, and a merged update.",
    },
    "testing": {
        "mission": "Add basic unit or UI tests to an app and learn assertions, test structure, and failure handling.",
        "output": "A project with passing tests for at least one component, function, or user flow.",
    },
    "figma": {
        "mission": "Create a simple wireframe and high-fidelity screen for the project before implementation.",
        "output": "A Figma file with one wireframe, one polished screen, and reusable design components.",
    },
    "python": {
        "mission": "Strengthen Python syntax, functions, and file handling through one practical automation or backend script.",
        "output": "A Python script or mini project with clean functions and documented usage.",
    },
    "django": {
        "mission": "Build one CRUD feature in Django to practice models, views, templates, forms, and routing.",
        "output": "A Django module with working create, read, update, and delete flow.",
    },
    "postgresql": {
        "mission": "Practice schema design, queries, and relational modeling by storing real app data in PostgreSQL.",
        "output": "A PostgreSQL-backed mini app with at least two related tables and basic queries.",
    },
    "docker": {
        "mission": "Containerize one app so it can run consistently with Docker locally or in deployment.",
        "output": "A Dockerfile and working container run command for your project.",
    },
    "machine learning": {
        "mission": "Learn the ML workflow by preparing data, training a small model, and evaluating the result.",
        "output": "A notebook or script showing dataset prep, model training, and evaluation metrics.",
    },
    "scikit-learn": {
        "mission": "Train a basic classification or regression model using scikit-learn and explain the pipeline clearly.",
        "output": "A scikit-learn project with preprocessing, model fitting, and evaluation summary.",
    },
    "tensorflow": {
        "mission": "Build a beginner neural-network project and understand model training, loss, and inference.",
        "output": "A TensorFlow notebook or script with one trained model and result explanation.",
    },
    "pytorch": {
        "mission": "Practice tensors, datasets, and training loops by building a small PyTorch model.",
        "output": "A PyTorch project with a working training loop and saved results.",
    },
    "fastapi": {
        "mission": "Create one API endpoint set with validation, routing, and documentation using FastAPI.",
        "output": "A FastAPI app with working endpoints and Swagger docs.",
    },
}


def build_learning_roadmap(missing_skills):
    roadmap = []
    for index, skill in enumerate(missing_skills[:6], start=1):
        plan = ROADMAP_PLAYBOOK.get(
            skill,
            {
                "mission": f"Study the fundamentals of {skill}, then apply it in one small practice project with clear scope.",
                "output": f"A portfolio-ready mini project or documented learning note for {skill}.",
            },
        )
        roadmap.append(
            {
                "week": f"Week {index}",
                "focus": skill,
                "mission": plan["mission"],
                "output": plan["output"],
            }
        )
    return roadmap


def build_interview_questions(role, matched_skills, missing_skills):
    return [
        f"Walk through a project where you used {matched_skills[0] if matched_skills else role['skills'][0]} to solve a real problem.",
        f"How would you prioritize learning {missing_skills[0] if missing_skills else role['skills'][1]} while still delivering team value?",
        f"Which tradeoffs matter most in {role['interview_themes'][0]} for a {role['title']} role?",
        "Describe how you collaborate with stakeholders when requirements are ambiguous.",
        "What metric would you track to prove your work improved the product or system?",
    ]
